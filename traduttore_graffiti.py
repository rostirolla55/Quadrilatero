import re
import sys
import os
from docx import Document
from googletrans import Translator

def process_document(input_path):
    if not os.path.exists(input_path):
        print(f"Errore: Il file {input_path} non esiste.")
        return

    directory = os.path.dirname(input_path)
    base_filename = os.path.basename(input_path)
    languages = ['en', 'es', 'fr']
    translator = Translator()

    for lang in languages:
        doc = Document(input_path)
        img_mapping = {}
        img_counter = 1
        
        print(f"\nLingua: [{lang}]")

        # --- FASE 1: MAPPATURA (Trova i nomi originali) ---
        for paragraph in doc.paragraphs:
            pattern = r"\[SPLIT_BLOCK:(.*?)\]"
            matches = re.findall(pattern, paragraph.text)
            for original_name in matches:
                if original_name not in img_mapping.values() and img_counter <= 5:
                    placeholder = f"##ID{img_counter}##" # Codice che il traduttore non tocca
                    img_mapping[placeholder] = original_name
                    img_counter += 1

        # --- FASE 2: PROTEZIONE E TRADUZIONE ---
        for paragraph in doc.paragraphs:
            # Sostituzione dei tag con placeholder prima della traduzione
            for placeholder, original_name in img_mapping.items():
                full_tag = f"[SPLIT_BLOCK:{original_name}]"
                if full_tag in paragraph.text:
                    for run in paragraph.runs:
                        if full_tag in run.text:
                            run.text = run.text.replace(full_tag, placeholder)

            # Traduzione dei run (SOLO se contengono testo e NON immagini/placeholder)
            for run in paragraph.runs:
                # Se il run contiene il placeholder o è vuoto o ha un'immagine, NON tradurre
                if not run.text.strip() or "##ID" in run.text:
                    continue
                
                # Protezione immagini: se il run ha elementi grafici XML, lo saltiamo
                if 'w:drawing' in run._element.xml or 'w:pict' in run._element.xml:
                    continue

                try:
                    translated = translator.translate(run.text, src='it', dest=lang).text
                    run.text = translated
                except Exception as e:
                    print(f"    Errore traduzione run: {e}")

        # --- FASE 3: RIPRISTINO INTEGRALE ---
        for paragraph in doc.paragraphs:
            # Rimuoviamo eventuali spazi folli che il traduttore mette vicino ai simboli #
            text_fix = paragraph.text.replace("# #", "##").replace("## ", "##").replace(" ##", "##")
            
            for placeholder, original_name in img_mapping.items():
                if placeholder in paragraph.text or placeholder in text_fix:
                    # Ricostruiamo il tag esattamente come richiesto
                    final_tag = f"[SPLIT_BLOCK:{original_name}]"
                    for run in paragraph.runs:
                        # Gestione flessibile per eventuali spazi inseriti dal traduttore nei simboli ##
                        run.text = run.text.replace(" # #", "##").replace("# # ", "##")
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, final_tag)

        # Salvataggio
        new_filename = base_filename.replace("_it.docx", f"_{lang}.docx")
        if new_filename == base_filename:
            new_filename = base_filename.replace(".docx", f"_{lang}.docx")
        
        output_file = os.path.join(directory, new_filename)
        doc.save(output_file)
        print(f" -> Salvato: {output_file}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Trascina un file .docx")
    else:
        process_document(sys.argv[1])